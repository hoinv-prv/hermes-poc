"""
convert_md_to_docx.py
Convert Markdown to DOCX using python-docx.

Requires:
    pip install python-docx

Supports:
  - Headings H1–H6
  - Paragraphs with bold (**), italic (*), inline code (`)
  - Bullet lists (-, *, +)
  - Numbered lists (1. 2. 3.)
  - Tables (| col | col |)
  - Code blocks (``` ... ```)
  - Horizontal rules (--- / ***)
  - HTML comments and <!-- Section N --> markers (stripped)

PDF output: open the .docx in Word → File > Export > PDF.
            On Windows with Word installed: `python -c "import subprocess; subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', 'file.docx'])"`
            Or use LibreOffice if available.

Usage:
    python convert_md_to_docx.py input.md [--out output.docx]
"""
import argparse
import re
import sys
from pathlib import Path


# ── Inline parsing ────────────────────────────────────────────────────────────

def parse_inline(text: str) -> list[tuple[str, dict]]:
    """Parse inline markdown → list of (text, styles) tuples."""
    # Strip links: [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    # Strip image refs: ![alt](url) → (nothing)
    text = re.sub(r"!\[[^\]]*\]\([^\)]*\)", "", text)

    result = []
    remaining = text
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")

    while remaining:
        m = pattern.search(remaining)
        if not m:
            if remaining:
                result.append((remaining, {}))
            break
        if m.start() > 0:
            result.append((remaining[: m.start()], {}))
        if m.group(1) is not None:
            result.append((m.group(1), {"bold": True, "italic": True}))
        elif m.group(2) is not None:
            result.append((m.group(2), {"bold": True}))
        elif m.group(3) is not None:
            result.append((m.group(3), {"italic": True}))
        elif m.group(4) is not None:
            result.append((m.group(4), {"mono": True}))
        remaining = remaining[m.end():]
    return result


def add_inline(para, text: str):
    """Add inline-formatted text to a paragraph."""
    from docx.shared import Pt

    for seg, styles in parse_inline(text):
        run = para.add_run(seg)
        if styles.get("bold"):
            run.bold = True
        if styles.get("italic"):
            run.italic = True
        if styles.get("mono"):
            run.font.name = "Courier New"
            run.font.size = Pt(9)


# ── Table helper ──────────────────────────────────────────────────────────────

def parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_separator_row(line: str) -> bool:
    return bool(re.match(r"^\|[-| :]+\|$", line.strip()))


def flush_table(doc, table_lines: list[str]):
    rows = []
    header_done = False
    for line in table_lines:
        if is_separator_row(line):
            header_done = True
            continue
        cells = parse_table_row(line)
        rows.append((cells, not header_done))

    if not rows:
        return

    n_cols = max(len(r[0]) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for i, (cells, is_header) in enumerate(rows):
        tr = table.rows[i]
        for j in range(n_cols):
            cell_text = cells[j] if j < len(cells) else ""
            cell = tr.cells[j]
            para = cell.paragraphs[0]
            para.clear()
            add_inline(para, cell_text)
            if is_header:
                for run in para.runs:
                    run.bold = True


# ── Main converter ────────────────────────────────────────────────────────────

def convert(input_path: Path, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    text = input_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    i = 0
    in_code = False
    in_comment = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_code():
        nonlocal in_code, code_lines
        if code_lines:
            para = doc.add_paragraph()
            para.style = "No Spacing"
            for cl in code_lines:
                run = para.add_run(cl + "\n")
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        code_lines = []
        in_code = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Multi-line HTML comment tracking ─────────────────────────────────
        if in_comment:
            if "-->" in stripped:
                in_comment = False
            i += 1
            continue

        # ── Code block ───────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if table_lines:
                flush_table(doc, table_lines)
                table_lines = []
            if in_code:
                flush_code()
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Table row ────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_lines.append(stripped)
            i += 1
            continue
        elif table_lines:
            flush_table(doc, table_lines)
            table_lines = []

        # ── HTML comment / section marker (strip) ────────────────────────────
        if stripped.startswith("<!--"):
            if "-->" not in stripped:
                in_comment = True
            i += 1
            continue

        # ── Image lines (any format — skip) ──────────────────────────────────
        if stripped.startswith("!["):
            i += 1
            continue

        # ── Heading ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            para = doc.add_paragraph(style=f"Heading {level}")
            add_inline(para, m.group(2))
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}$", stripped):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────────────
        m = re.match(r"^[-*+]\s+(.+)$", stripped)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, m.group(1))
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────────────
        m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_inline(para, m.group(1))
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        para = doc.add_paragraph()
        add_inline(para, stripped)
        i += 1

    # Flush any unclosed blocks
    if in_code:
        flush_code()
    if table_lines:
        flush_table(doc, table_lines)

    doc.save(str(output_path))
    print(f"Saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX")
    parser.add_argument("input", help="Input .md file")
    parser.add_argument("--out", help="Output .docx path (default: same dir, .docx extension)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.out) if args.out else input_path.with_suffix(".docx")
    print(f"Converting {input_path.name} → {output_path.name}...")
    convert(input_path, output_path)


if __name__ == "__main__":
    main()
